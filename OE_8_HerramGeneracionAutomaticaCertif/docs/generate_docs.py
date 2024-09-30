# docs/generate_docs.py
from docx import Document
import pandas as pd
import locale
from docs.replace_data import *
from utils.helpers import set_und  # Asegúrate de que set_und esté correctamente definido en helpers.py
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

def create_certificate(template_path, output_path, data):
    doc = Document(template_path)
    # Lógica para llenar la plantilla con datos
    doc.save(output_path)
    
def create_certificate_no_station(nombres, apellidos, correo, descrip_solicit, clickinfo):
    template_path = "PlantillaOficioLamentoSinEstaciones.docx"
    doc = Document(template_path)
    doc = reemplazar_datos_noEMC(doc, nombres, apellidos, correo, descrip_solicit, clickinfo)
    return doc

def create_certificate_nodata(nombres, apellidos, correo, dias, meses, ano, selected_variable, estacion_seleccionada, descrip_solicit):
    template_path = "PlantillaOficioLamentoSinDatos.docx"
    doc = Document(template_path)
    doc = reemplazar_datos_nan(doc, nombres, apellidos, correo, dias, meses, ano, selected_variable, estacion_seleccionada, descrip_solicit)
    return doc

def select_plantilla(selected_variable, stationdf_fnl, nombres, apellidos, correo, dias, meses, ano, estacion_seleccionada, descrip_solicit):
    plantillas_por_variable = {
        "Precipitación total diaria": "PlantillaPrecipDiaria.docx",
        "Precipitación total mensual": "PlantillaPrecipMensual.docx",
        "Precipitación total anual": "PlantillaPrecipAnual.docx",
        "Temperatura máxima diaria": "PlantillaTemperatDiaria.docx",
        "Temperatura máxima media mensual": "PlantillaTemperatMensual.docx",
        "Temperatura máxima media anual": "PlantillaTemperatAnual.docx",
        "Temperatura mínima diaria": "PlantillaTemperatDiaria.docx",
        "Temperatura mínima media mensual": "PlantillaTemperatMensual.docx",
        "Temperatura mínima media anual": "PlantillaTemperatAnual.docx",
        "Temperatura del aire media diaria": "PlantillaTemperatDiaria.docx",
        "Temperatura del aire media mensual": "PlantillaTemperatMensual.docx",
        "Temperatura del aire media anual": "PlantillaTemperatAnual.docx",
        "Velocidad del viento horaria": "PlantillaVelVientoHoraria.docx",
        "Velocidad del viento media diaria": "PlantillaVelVientoDiaria.docx",
        "Velocidad del viento media mensual": "PlantillaVelVientoMensual.docx",
        "Velocidad del viento media anual": "PlantillaVelVientoAnual.docx"
    }
    
    if selected_variable not in plantillas_por_variable:
        raise ValueError("Variable seleccionada no tiene plantilla asociada")

    nombre_plantilla = plantillas_por_variable[selected_variable]
    doc = Document(nombre_plantilla)
    
    # Determina si el DataFrame está vacío para seleccionar la plantilla
    clave_plantilla = "Sin Datos" if stationdf_fnl.empty else selected_variable
    if clave_plantilla == "Sin Datos":
        doc = create_certificate_nodata(nombres, apellidos, correo, dias, meses, ano, selected_variable, estacion_seleccionada, descrip_solicit)
    else:
        clave_plantilla = selected_variable
        locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
        stationdf_fnl['Fecha'] = pd.to_datetime(stationdf_fnl['Fecha'], format='%Y-%m-%d %H:%M:%S.%f')
        sufij_periodo = {"horaria": '%Y-%m-%d %H:%M', "diaria": '%Y-%m-%d', "mensual": '%Y-%m', "anual": '%Y'}
        for sufij, period in sufij_periodo.items():
            if selected_variable.endswith(sufij):
                stationdf_fnl['Fecha'] = stationdf_fnl['Fecha'].dt.strftime(period)

        if clave_plantilla == "Precipitación total diaria":
            primer_fecha = pd.to_datetime(stationdf_fnl['Fecha'].iloc[0])
            dia = primer_fecha.day
            mes_nm = primer_fecha.strftime('%B')
            ano_p = primer_fecha.year
            primer_valor = stationdf_fnl['Valor'].iloc[0]

            doc = reemplazar_datos_precip(doc, nombres, apellidos, dias, meses, ano, selected_variable, estacion_seleccionada, dia, mes_nm, ano_p, primer_valor)
        else:
            doc = reemplazar_datos(doc, nombres, apellidos, dias, meses, ano, selected_variable, estacion_seleccionada)

        und = set_und(selected_variable)
        stationdf_fnl.loc[:, 'Valor'] = stationdf_fnl['Valor'].fillna('ND')
        stationdf_fnl = stationdf_fnl.rename(columns={'Valor': f"{selected_variable} ({und})"})
    return doc, nombre_plantilla

def insert_table_in_doc(doc, data_frame, selected_variable):
    table = doc.add_table(rows=data_frame.shape[0] + 1, cols=data_frame.shape[1])
    #try:
    #    table.style = 'Table Grid'
    #except (KeyError, ValueError):
    #    try:
    #        table.style = 'Light Grid'
    #    except (KeyError, ValueError):
            # Si no se encuentra un estilo adecuado, se deja la tabla sin estilo específico
    #        pass

    # Aplicar bordes manualmente
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            tcPr.append(tcBorders)

            # Aplicar fuente Verdana y tamaño de 11 puntos
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Verdana'
                    run.font.size = Pt(11)

    # Poner el encabezado de la tabla en negrita y aplicar fuente Verdana
    for i, column in enumerate(data_frame.columns):
        cell = table.cell(0, i)
        cell.text = column
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Verdana'
            run.font.size = Pt(11)

    # Rellenar las celdas con los datos
    for i, column in enumerate(data_frame.columns):
        for j, value in enumerate(data_frame[column]):
            cell = table.cell(j + 1, i)
            cell.text = str(value)
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Verdana'
                run.font.size = Pt(11)
    #for i, column in enumerate(data_frame.columns):
    #    table.cell(0, i).text = column
    #    for j, value in enumerate(data_frame[column]):
    #        table.cell(j + 1, i).text = str(value)

    def move_table_after(table, paragraph):
        tbl, p = table._tbl, paragraph._p
        p.addnext(tbl)

    marcador_texto = "Datos disponibles en la base de datos IDEAM:"
    for paragraph in doc.paragraphs:
        if marcador_texto in paragraph.text:
            doc.add_paragraph('\n')
            move_table_after(table, paragraph)
            break

    row = table.rows[0]
    tr = row._tr
    tblPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    tblPr.append(tblHeader)
    
    return doc


