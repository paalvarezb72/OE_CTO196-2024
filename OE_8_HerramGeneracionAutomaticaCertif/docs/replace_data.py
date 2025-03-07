# docs/replace_data.py
from datetime import datetime
from docx.shared import Pt

def reemplazar_datos_en_runs(paragraph, search_text, replace_text):
    text = paragraph.text
    if search_text in text:
        new_text = text.replace(search_text, replace_text)
        for run in paragraph.runs:
            run.text = ''
            run.font.name = 'Verdana'
            run.font.size = Pt(11)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

def reemplazar_datos_noEMC(tpers, doc, nombres, apellidos, correo, descrip_solicit, clickinfo):#, lat_pi, lon_pi):
    if tpers == 'Persona jurídica':
        datos = {
            "{{NOMBRES}}": nombres.upper(),
            "{{CORREO}}": correo,
            "{{DESCRIP_SOLICIT}}": descrip_solicit,
            "{{LAT_PI}}": str(clickinfo[0]),
            "{{LONG_PI}}": str(clickinfo[1])
        }
    elif tpers == 'Persona natural':
        datos = {
            "{{NOMBRES}}": nombres.upper(),
            "{{APELLIDOS}}": apellidos.upper(),
            "{{CORREO}}": correo,
            "{{DESCRIP_SOLICIT}}": descrip_solicit,
            "{{LAT_PI}}": str(clickinfo[0]),
            "{{LONG_PI}}": str(clickinfo[1])
        }
        
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)
    return doc

def reemplazar_datos_nan(tpers, doc, nombres, apellidos, correo, finicio, ffin, selected_variable, estacion_seleccionada, descrip_solicit):
    if not finicio or not ffin:
        return "Selecciona ambas fechas."
    
    # Convertir las fechas en objetos datetime
    fecha_inicio = datetime.fromisoformat(finicio)
    fecha_fin = datetime.fromisoformat(ffin)

    # Inicialización de textos
    dias = ""
    meses = ""
    anos = ""

    # Verificar la variable seleccionada: Anual, Mensual o Diaria
    if 'anual' in selected_variable.lower():
        # Solo interesa el año o años
        if fecha_inicio.year == fecha_fin.year:
            anos = str(fecha_inicio.year)
        else:
            anos = f"{fecha_inicio.year} a {fecha_fin.year}"

    elif 'mensual' in selected_variable.lower():
        # Se consideran meses y años
        if fecha_inicio.year == fecha_fin.year:
            if fecha_inicio.month == fecha_fin.month:
                meses = fecha_inicio.strftime('%B')
            else:
                meses = f"{fecha_inicio.strftime('%B')} a {fecha_fin.strftime('%B')} de"
            anos = str(fecha_inicio.year)
        else:
            meses = f"{fecha_inicio.strftime('%B')} de {fecha_inicio.year} a {fecha_fin.strftime('%B')} de {fecha_fin.year}"
            anos = ""

    elif 'diaria' or 'horaria' in selected_variable.lower():
        # Se consideran días, meses y años
        if fecha_inicio.year == fecha_fin.year and fecha_inicio.month == fecha_fin.month:
            # Mismo año y mismo mes
            dias = f"{fecha_inicio.day} al {fecha_fin.day} de"
            meses = f"{fecha_inicio.strftime('%B')} de"
            anos = str(fecha_inicio.year)
        elif fecha_inicio.year == fecha_fin.year:
            # Mismo año pero diferente mes
            dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de"
            meses = ""  # No necesitamos repetir el mes aquí
            anos = str(fecha_inicio.year)
        else:
            # Diferente año
            dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} de {fecha_inicio.year} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de {fecha_fin.year}"
            meses = ""  # No necesitamos repetir el mes aquí
            anos = ""  # No necesitamos repetir el año aquí
    
    if tpers == 'Persona natural':
        apellidosm = apellidos.upper()
    elif tpers == 'Persona jurídica':
        apellidosm = " "

    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidosm,
        "{{CORREO}}": correo,
        "{{DIAS}}": dias,
        "{{MESES}}": meses,
        "{{AÑO}}": anos,
        "{{VARIABLE}}": selected_variable,
        "{{ESTACION}}": estacion_seleccionada['nombre'],
        "{{DESCRIP_SOLICIT}}": descrip_solicit
    }
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)
    return doc

def reemplazar_datos_precipd(tpers, doc, nombres, apellidos, finicio, ffin, selected_variable, estacion_seleccionada, dia, mes_nm, ano_p, primer_valor):
    if not finicio or not ffin:
        return "Selecciona ambas fechas."
    
    # Convertir las fechas en objetos datetime
    fecha_inicio = datetime.fromisoformat(finicio)
    fecha_fin = datetime.fromisoformat(ffin)

    # Inicialización de textos
    dias = ""
    meses = ""
    anos = ""

    # Verificar la variable seleccionada: Anual, Mensual o Diaria
    if 'anual' in selected_variable.lower():
        # Solo interesa el año o años
        if fecha_inicio.year == fecha_fin.year:
            anos = str(fecha_inicio.year)
        else:
            anos = f"{fecha_inicio.year} a {fecha_fin.year}"

    elif 'mensual' in selected_variable.lower():
        # Se consideran meses y años
        if fecha_inicio.year == fecha_fin.year:
            if fecha_inicio.month == fecha_fin.month:
                meses = fecha_inicio.strftime('%B')
            else:
                meses = f"{fecha_inicio.strftime('%B')} a {fecha_fin.strftime('%B')} de"
            anos = str(fecha_inicio.year)
        else:
            meses = f"{fecha_inicio.strftime('%B')} de {fecha_inicio.year} a {fecha_fin.strftime('%B')} de {fecha_fin.year}"
            anos = ""

    elif 'diaria' or 'horaria' in selected_variable.lower():
        # Se consideran días, meses y años
        if fecha_inicio.year == fecha_fin.year and fecha_inicio.month == fecha_fin.month:
            # Mismo año y mismo mes
            dias = f"{fecha_inicio.day} al {fecha_fin.day} de"
            meses = f"{fecha_inicio.strftime('%B')} de"
            anos = str(fecha_inicio.year)
        elif fecha_inicio.year == fecha_fin.year:
            # Mismo año pero diferente mes
            dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de"
            meses = ""  # No necesitamos repetir el mes aquí
            anos = str(fecha_inicio.year)
        else:
            # Diferente año
            dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} de {fecha_inicio.year} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de {fecha_fin.year}"
            meses = ""  # No necesitamos repetir el mes aquí
            anos = ""  # No necesitamos repetir el año aquí
    
    if tpers == 'Persona natural':
        apellidosm = apellidos.upper()
    elif tpers == 'Persona jurídica':
        apellidosm = " "

    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidosm,
        "{{DIAS}}": dias,
        "{{MESES}}": meses,
        "{{AÑO}}": anos,
        "{{VARIABLE}}": selected_variable,
        "{{ESTACION}}": estacion_seleccionada['nombre'],
        "{{LATITUD}}": str(estacion_seleccionada['latitud']),
        "{{LONGITUD}}": str(estacion_seleccionada['longitud']),
        "{{ALTITUD}}": str(estacion_seleccionada['altitud']),
        "{{MUNICIPIO}}": str(estacion_seleccionada['MUNICIPIO']),
        "{{DEPARTAMENTO}}": str(estacion_seleccionada['DEPARTAMENTO']),
        "{{DIA}}": str(dia),
        "{{MES}}": str(mes_nm),
        "{{AÑO_P}}": str(ano_p),
        "{{PRIMER_DATO}}": str(primer_valor),
        "{{PRIMER_DATO_HA}}": str(primer_valor * 10),
    }

    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)

    return doc

# def reemplazar_datos_precipd(tpers, doc, nombres, apellidos, finicio, ffin, selected_variable, estacion_seleccionada, dia, mes_nm, ano_p, primer_valor):
#     if not finicio or not ffin:
#         return "Selecciona ambas fechas."
    
#     # Convertir las fechas en objetos datetime
#     fecha_inicio = datetime.fromisoformat(finicio)
#     fecha_fin = datetime.fromisoformat(ffin)

#     # Inicialización de textos
#     dias = ""
#     meses = ""
#     anos = ""

#     # Verificar la variable seleccionada: Anual, Mensual o Diaria
#     if 'anual' in selected_variable.lower():
#         # Solo interesa el año o años
#         if fecha_inicio.year == fecha_fin.year:
#             anos = str(fecha_inicio.year)
#         else:
#             anos = f"{fecha_inicio.year} a {fecha_fin.year}"

#     elif 'mensual' in selected_variable.lower():
#         # Se consideran meses y años
#         if fecha_inicio.year == fecha_fin.year:
#             if fecha_inicio.month == fecha_fin.month:
#                 meses = fecha_inicio.strftime('%B')
#             else:
#                 meses = f"{fecha_inicio.strftime('%B')} a {fecha_fin.strftime('%B')} de"
#             anos = str(fecha_inicio.year)
#         else:
#             meses = f"{fecha_inicio.strftime('%B')} de {fecha_inicio.year} a {fecha_fin.strftime('%B')} de {fecha_fin.year}"
#             anos = ""

#     elif 'diaria' in selected_variable.lower():
#         # Se consideran días, meses y años
#         if fecha_inicio.year == fecha_fin.year and fecha_inicio.month == fecha_fin.month:
#             dias = f"{fecha_inicio.day} al {fecha_fin.day}"
#             meses = fecha_inicio.strftime('%B')
#             anos = str(fecha_inicio.year)
#         elif fecha_inicio.year == fecha_fin.year:
#             dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} al {fecha_fin.day} de {fecha_fin.strftime('%B')}"
#             anos = str(fecha_inicio.year)
#         else:
#             dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} de {fecha_inicio.year} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de {fecha_fin.year}"
    

    
#     if tpers == 'Persona natural':
#         apellidosm = apellidos.upper()
#     elif tpers == 'Persona jurídica':
#         apellidosm = " "

#     datos = {
#         "{{NOMBRES}}": nombres.upper(),
#         "{{APELLIDOS}}": apellidosm,
#         "{{DIAS}}": dias,
#         "{{MESES}}": meses,
#         "{{AÑO}}": anos,
#         "{{VARIABLE}}": selected_variable,
#         "{{ESTACION}}": estacion_seleccionada['nombre'],
#         "{{LATITUD}}": str(estacion_seleccionada['latitud']),
#         "{{LONGITUD}}": str(estacion_seleccionada['longitud']),
#         "{{ALTITUD}}": str(estacion_seleccionada['altitud']),
#         "{{MUNICIPIO}}": str(estacion_seleccionada['MUNICIPIO']),
#         "{{DEPARTAMENTO}}": str(estacion_seleccionada['DEPARTAMENTO']),
#         "{{DIA}}": str(dia),
#         "{{MES}}": str(mes_nm),
#         "{{AÑO_P}}": str(ano_p),
#         "{{PRIMER_DATO}}": str(primer_valor),
#         "{{PRIMER_DATO_HA}}": str(primer_valor * 10),
#     }

#     for p in doc.paragraphs:
#         for key, value in datos.items():
#             reemplazar_datos_en_runs(p, key, value)

#     return doc

def reemplazar_datos(tpers, doc, nombres, apellidos, finicio, ffin, selected_variable, estacion_seleccionada):
    if not finicio or not ffin:
        return "Selecciona ambas fechas."
    
    # Convertir las fechas en objetos datetime
    fecha_inicio = datetime.fromisoformat(finicio)
    fecha_fin = datetime.fromisoformat(ffin)

    # Inicialización de textos
    dias = ""
    meses = ""
    anos = ""

    # Verificar la variable seleccionada: Anual, Mensual o Diaria
    if 'anual' in selected_variable.lower():
        # Solo interesa el año o años
        if fecha_inicio.year == fecha_fin.year:
            anos = str(fecha_inicio.year)
        else:
            anos = f"{fecha_inicio.year} a {fecha_fin.year}"

    elif 'mensual' in selected_variable.lower():
        # Se consideran meses y años
        if fecha_inicio.year == fecha_fin.year:
            if fecha_inicio.month == fecha_fin.month:
                meses = fecha_inicio.strftime('%B')
            else:
                meses = f"{fecha_inicio.strftime('%B')} a {fecha_fin.strftime('%B')} de"
            anos = str(fecha_inicio.year)
        else:
            meses = f"{fecha_inicio.strftime('%B')} de {fecha_inicio.year} a {fecha_fin.strftime('%B')} de {fecha_fin.year}"
            anos = ""

    elif 'diaria' or 'horaria' in selected_variable.lower():
        # Se consideran días, meses y años
        if fecha_inicio.year == fecha_fin.year and fecha_inicio.month == fecha_fin.month:
            # Mismo año y mismo mes
            dias = f"{fecha_inicio.day} al {fecha_fin.day} de"
            meses = f"{fecha_inicio.strftime('%B')} de"
            anos = str(fecha_inicio.year)
        elif fecha_inicio.year == fecha_fin.year:
            # Mismo año pero diferente mes
            dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de"
            meses = ""  # No necesitamos repetir el mes aquí
            anos = str(fecha_inicio.year)
        else:
            # Diferente año
            dias = f"{fecha_inicio.day} de {fecha_inicio.strftime('%B')} de {fecha_inicio.year} al {fecha_fin.day} de {fecha_fin.strftime('%B')} de {fecha_fin.year}"
            meses = ""  # No necesitamos repetir el mes aquí
            anos = ""  # No necesitamos repetir el año aquí

    if tpers == 'Persona natural':
        apellidosm = apellidos.upper()
    elif tpers == 'Persona jurídica':
        apellidosm = " "

    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidosm,
        "{{DIAS}}": dias,  # Manejar caso donde dias es None
        "{{MESES}}": meses,  # Manejar caso donde meses es None
        "{{AÑO}}": anos, #if ano else "",  # Manejar caso donde ano es None
        "{{VARIABLE}}": selected_variable,
        "{{ESTACION}}": estacion_seleccionada['nombre'],
        "{{LATITUD}}": str(estacion_seleccionada['latitud']),
        "{{LONGITUD}}": str(estacion_seleccionada['longitud']),
        "{{ALTITUD}}": str(estacion_seleccionada['altitud']),
        "{{MUNICIPIO}}": str(estacion_seleccionada['MUNICIPIO']),
        "{{DEPARTAMENTO}}": str(estacion_seleccionada['DEPARTAMENTO']),
    }
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)
    return doc
