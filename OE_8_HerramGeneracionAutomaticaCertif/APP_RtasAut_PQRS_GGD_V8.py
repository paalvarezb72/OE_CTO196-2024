import pandas as pd
import dash
import plotly.express as px
import prestodb
import data.derivadasclc as dvd
import locale
import traceback
import re
import numpy as np
import json
import geopy.distance
from dash import html
from datetime import datetime
from dash.dependencies import Input, Output, State
from dash import dcc
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# Lectura de CNE de convencionales activas
#EstCON_Act = pd.read_excel('EstacCON_IDEAM_Activas.xls')
EstCON_Act = pd.read_table('EstacCON_IDEAM_Activ_AltLimSupRP.txt', sep=';')

# Se genera usuario prestodb para la extracción de datos
conn = prestodb.dbapi.connect(
    host='172.16.50.20',
    port=8080,
    user='Paola',
    catalog='raw',
    schema='cassandra',
)
cur = conn.cursor()

# Se crea la aplicación Dash
app = dash.Dash(__name__)

app.title = "Certificaciones estado tiempo y clima"

fig = px.scatter_mapbox(EstCON_Act,
                        lat="latitud",
                        lon="longitud",
                        hover_name="nombre",  # Nombre a mostrar en el hover
                        hover_data={"altitud": True},
                        zoom=10,
                        mapbox_style="open-street-map")

# Se define la estructura de la página web
app.layout = html.Div([
    # Banner
    html.Div(
        id="banner",
        className="banner",
        children=[
            # Contenedor para la imagen superior
            html.Div([
                html.Img(src=app.get_asset_url('Banner_GovCo.png'), style={
                         'width': '100%', 'display': 'block', 'margin-top': '0', 'padding-top': '0'}),
            ], style={'width': '100%', 'display': 'block', 'margin-top': '0', 'padding-top': '0'}),

            # Espaciador
            html.Div(style={'height': '20px', 'width': '100%'}),

            # Contenedor para las otras dos imágenes
            html.Div([
                html.Img(src=app.get_asset_url('Ideam_Colombia_logo2.png'), style={
                         'display': 'inline-block', 'width': 'auto'}),
                html.Img(src=app.get_asset_url('Colombia_Potencia_Vida_Ambiente2.png'), style={
                         'display': 'inline-block', 'width': '27%'}),
            ], style={'width': '100%', 'display': 'flex', 'justify-content': 'space-between', 'align-items': 'center'}),
        ]
    ),

    # Contenedor del resto de contenidos
    html.Div(
        # Ajusta los márgenes según necesites
        style={'margin-left': '10%', 'margin-right': '10%'},
        children=[
            html.H1("Certificaciones del estado del tiempo y del clima - IDEAM", style={'font-family': 'arial', 'text-align': 'center'},
                    title="Datos obtenidos de la red IDEAM "),

            html.P("Respetado usuario, diligiencie todos los campos para obtener su certificación:",
                   style={'font-family': 'arial', 'text-align': 'justify', 'margin-bottom': '20px'}),

            html.Div([
                html.Div([
                    html.Label("Nombre(s):", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Input(id="nombres-input", type="text"),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),

                html.Div([
                    html.Label("Apellido(s):", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Input(id="apellidos-input", type="text"),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),

                html.Div([
                    html.Label("Correo:", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Input(id="correo-input", type="text"),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),
            ], style={'display': 'flex', 'justify-content': 'space-between', 'flex-wrap': 'wrap'}),

            # Espaciador
            html.Div(style={'height': '10px', 'width': '100%'}),

            html.P("Por favor, escoja las fechas de interés para su certificación. Si  requiere\
                   un rango, seleccione solo los días/meses/años correspondientes a la fecha inicial y final:",
                   style={'font-family': 'arial', 'text-align': 'justify', 'margin-bottom': '20px'}),

            html.Div([
                html.Div([
                    html.Label("Día(s):", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Dropdown(
                        id="dias-dropdown",
                        options=[{"label": str(i), "value": i}
                                 for i in range(1, 32)],
                        multi=True,
                        # Ajusta el ancho del dropdown al 100% del contenedor
                        style={'width': '100%', 'font-family': 'arial'}
                    ),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),  # Ajusta el padding para espacio interno y un min-width

                html.Div([
                    html.Label("Mes(es):", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Dropdown(
                        id="meses-dropdown",
                        options=[
                            {"label": "Enero", "value": "enero"}, {
                                "label": "Febrero", "value": "febrero"},
                            {"label": "Marzo", "value": "marzo"}, {
                                "label": "Abril", "value": "abril"},
                            {"label": "Mayo", "value": "mayo"}, {
                                "label": "Junio", "value": "junio"},
                            {"label": "Julio", "value": "julio"}, {
                                "label": "Agosto", "value": "agosto"},
                            {"label": "Septiembre", "value": "septiembre"}, {
                                "label": "Octubre", "value": "octubre"},
                            {"label": "Noviembre", "value": "noviembre"}, {
                                "label": "Diciembre", "value": "diciembre"}
                        ],
                        multi=True,
                        # Igual aquí
                        style={'width': '100%', 'font-family': 'arial'}
                    ),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),  # Y aquí

                html.Div([
                    html.Label("Año(s):", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Dropdown(
                        id="ano-dropdown",
                        options=[{"label": str(i), "value": i}
                                 for i in range(2024, 1949, -1)],
                        multi=True,
                        # Y finalmente aquí
                        style={'width': '100%', 'font-family': 'arial'}
                    ),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'})  # Ajusta estos valores según necesites
            ], style={'display': 'flex', 'justify-content': 'space-between', 'flex-wrap': 'wrap'}),  # Mejora la disposición de los elementos

            # Espaciador
            html.Div(
                html.P("Por favor, elija la variable meteorológica y su periodicidad de interés:",
                       style={'font-family': 'arial', 'text-align': 'justify', 'margin-bottom': '20px'}),
                style={'height': '20px', 'width': '100%'}),

            # Espaciador
            html.Div(style={'height': '20px', 'width': '100%'}),

            # Sección de escoger la variable y su temporalidad
            html.Div([
                html.Div([
                    html.Label("Variable meteorológica:", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Dropdown(
                        id="variable-dropdown",
                        options=[
                            {"label": "Precipitación", "value": "Precipitación"},
                            {"label": "Temperatura máxima",
                                "value": "Temperatura máxima"},
                            {"label": "Temperatura mínima",
                                "value": "Temperatura mínima"},
                            {"label": "Temperatura del aire",
                                "value": "Temperatura del aire"},
                            {"label": "Velocidad del viento", "value": "Velocidad del viento"}],
                        style={'font-family': 'arial'}
                    ),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),

                html.Div([
                    html.Label("Variable y periodo:", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Dropdown(id="tiposerie-dropdown",
                                 style={'font-family': 'arial'})
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '300px'}),

            ], style={'display': 'flex', 'justify-content': 'space-between', 'flex-wrap': 'wrap'}),

            # Espaciador
            html.Div(style={'height': '20px', 'width': '100%'}),

            html.P("En el siguiente mapa, busque por favor su área de interés, \
                   desplácese en este manteniendo oprimido el clic izquierdo del cursor \
                   y observe el nombre de la estación meteorológica más cercana, tenga en \
                   cuenta el nombre obtenido.",
                   style={'font-family': 'arial', 'text-align': 'justify', 'margin-bottom': '20px'}),

            html.P("Recuerde que una estación se considera cercana si se encuentra a una distancia\
                   inferior a 10 km planos del punto de interés -si la variable de interés es velocidad\
                    del viento, la distancia debe ser inferior a 3 km-, no obstante, considere que aún\
                    dentro de ese radio, una diferencia elevada de altitudes de la estación con respecto\
                    a su punto de interés, es decir, superior a los 200msnm, (entre otros factores) \
                    puede significar que esta no sea representativa para su zona de estudio. \
                    A continuación, puede calcular las distancias lineales siguiendo las instrucciones.",
                   style={'font-family': 'arial', 'text-align': 'justify', 'margin-bottom': '20px', 'font-style': 'italic'}),

            html.Div([

                html.P("Calcule la distancia de la estación cercana a su punto de interés",
                       style={'font-size': 16, 'font-family': 'arial', 'margin-top': '20px', 'font-weight': 'bold'}),
                html.Div(id="distance-info",
                         style={'height': '20px', 'width': '100%'}),

                # Espaciador
                html.Div(style={'height': '20px', 'width': '100%'}),

                html.Div([
                    html.Label("Latitud del punto de interés:", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Input(id="lat-input", type="number",
                              placeholder="P. e.: -2.64"),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '200px'}),

                html.Div([
                    html.Label("Longitud del punto de interés:", style={
                               'font-family': 'arial', 'color': '#5D5D5D', 'font-size': 15}),
                    dcc.Input(id="lon-input", type="number",
                              placeholder="P. e.: -74.86"),
                ], style={'flex': 1, 'padding': '0 10px', 'min-width': '200px'}),

                html.Button("Calcular distancia", id="calculate-distance-btn",
                            style={'font-family': 'arial', 'flex': 1, 'padding': '0 10px', 'min-width': '150px'}),
                html.Button('Reiniciar selección', id='reset-button', n_clicks=0,
                            style={'font-family': 'arial', 'flex': 1, 'padding': '0 10px', 'min-width': '150px'}),
            ], style={'display': 'flex', 'justify-content': 'space-between', 'flex-wrap': 'wrap'}),

            dcc.Graph(id="mapa-estaciones",
                      figure=fig,
                      config={"scrollZoom": True, "responsive": True},
                      style={"height": "500px", 'font-family': 'arial'},
            ),
            
            html.Div([
                html.P("Si ninguna estación es cercana o de similar altitud a su punto de interés,\
                       por favor seleccione el siguiente ítem (recuerde antes haber registrado\
                       las coordenadas de este con la anterior herramienta)",
                       style={'font-family': 'arial', 'font-size': 13}),#, 'margin-top': '20px', }),
                
                dcc.RadioItems(
                    options=[{'label': 'Sin estaciones representativas', 'value': 'no_station'}],
                    value=None,
                    id='sin-estaciones',
                    style={'font-family': 'arial', 'font-size': 13}
                ),
                html.Div(id='hidden-div', style={'display': 'none'})
                
            ], style={'display': 'flex', 'justify-content':'space-between'}),#, 'flex-wrap': 'wrap'}),

            html.P("A partir del ejercicio anterior, seleccione la estación de su interés",
                   style={'font-family': 'arial', 'margin-top': '20px'}),

            dcc.Dropdown(
                id='estacion-dropdown',
                options=[{'label': nombre, 'value': nombre}
                         for nombre in EstCON_Act['nombre']],
                style={'font-family': 'arial'},
                value=EstCON_Act['nombre'][0]  # Valor por defecto
            ),

            # Espaciador
            html.Div(style={'height': '20px', 'width': '100%'}),

            html.Div(id='output-state'),

            html.Div(style={'height': '10px', 'width': '100%'}),

            html.Button("Generar Certificación", style={
                        'font-family': 'arial'}, id="generar-button"),

            html.P("Si requiere una certificación con información adicional -fenómenos, anomalías,\
                   análisis climatológicos, entre otros-, por favor tramite su solicitud al correo\
                       contacto@ideam.gov.co",
                   style={'font-family': 'arial', 'text-align': 'justify', 'margin-bottom': '20px'}),

            # Espaciador
            html.Div(style={'height': '20px', 'width': '100%'})

        ])
], style={'width': '100%', 'max-width': '1200px', 'margin': '0 auto'})


def construir_rango_fechas(dias, meses, ano):
    # Asegurarse de que los valores sean convertidos a enteros
    dias = list(map(int, dias))
    meses_dict = {"enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
                  "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12}
    meses = list(map(lambda mes: meses_dict[mes], meses))
    ano = list(map(int, ano))

    # Se establece la fecha como rango desde la inicial hasta la final
    fecha_inicio = datetime(ano[0], meses[0], dias[0]
                            ).strftime('%Y-%m-%d %H:%M:%S')
    fecha_fin = datetime(ano[-1], meses[-1], dias[-1]
                         ).strftime('%Y-%m-%d %H:%M:%S')
    return fecha_inicio, fecha_fin


def obtener_sensor(selected_variable):
    # Diccionario que mapea prefijos de variables a códigos de sensor
    prefijo_a_sensor = {
        "Precipitación": 'PTPM_CON',
        "Temperatura máxima": 'TMX_CON',
        "Temperatura mínima": 'TMN_CON',
        "Temperatura del aire": 'TSSM_CON',
        "Velocidad del viento": 'VVAG_CON'
    }

    # Itera sobre los prefijos en el diccionario
    for prefijo, sensor in prefijo_a_sensor.items():
        if selected_variable.startswith(prefijo):
            return sensor

    # Retorna None o un valor por defecto si no se encuentra un prefijo correspondiente
    return None


def set_und(selected_variable):
    # Diccionario que mapea prefijos de variables a códigos de sensor
    prefijo_a_und = {
        "Precipitación": 'mm',
        "Temperatura": '°C',
        "Velocidad del viento": 'km/h'
    }

    # Itera sobre los prefijos en el diccionario
    for prefijo, und in prefijo_a_und.items():
        if selected_variable.startswith(prefijo):
            return und

    # Retorna None o un valor por defecto si no se encuentra un prefijo correspondiente
    return None


def construir_codestacion(estacion_seleccionada):
    # Se construye código estación según se solicita en Cassandra
    cod_est = str(estacion_seleccionada['CODIGO'])
    if len(cod_est) == 8:
        # Agregar dos ceros a la izquierda si el código tiene 8 dígitos
        codigo_cass = cod_est.zfill(10)
    elif len(cod_est) == 10:
        # Dejar el código tal cual si tiene 10 dígitos
        codigo_cass = cod_est
    else:
        # Manejar otros casos según sea necesario
        codigo_cass = "Código no válido"

    return codigo_cass


def modifdfprecip_ClasifLimSup(df, EstCON_Act, selected_variable, codestacion):
    # Definir los bins y las etiquetas para la clasificación
    bins = [0, 0.1, 10, 20, 40, 60, float('inf')]
    labels = ['Tiempo seco', 'Lluvia ligera', 'Lluvia ligera a moderada',
              'Lluvia moderada a fuerte', 'Lluvia fuerte a torrencial', 'Lluvia Torrencial']

    if selected_variable == "Precipitación total diaria":
        # Volver numérico 'codestacion
        codestacion_n = re.sub("[^0-9]", "", codestacion)
        try:
            codestacion_n = int(codestacion_n)
        except ValueError:
            print("codestacion contiene caracteres no convertibles a int.")
        # Obtener el límite superior para la estación seleccionada
        limite_superior = EstCON_Act.loc[EstCON_Act['CODIGO'] == codestacion_n, 'LimSup']#.values[0]
        if not limite_superior.empty:
            limite_superior = limite_superior.values[0]
        else:
            # Manejo del caso en que no se encuentre un limite_superior
            print(f"No se encontró un límite superior para el código de estación {codestacion}.")
            df['Valor'] = 'ND'
            df['Clasificación'] = 'No aplica'
            return df
        
        # Evaluar cada fila para determinar si el valor excede el límite superior
        def evaluar_fila(row):
            if row['Valor'] > limite_superior:
                return (np.nan, 'No aplica')  # Asignar NaN a 'Valor' y 'No aplica' a 'Clasificación'
            else:
                return (row['Valor'], pd.cut([row['Valor']], bins=bins, labels=labels, right=False)[0])

        # Aplicar la función evaluar_fila y separar los resultados en dos columnas
        df[['Valor', 'Clasificación']] = df.apply(evaluar_fila, axis=1, result_type='expand')

    return df


def construir_descripsolicit(var_periodo):
    descripcion = "Solicitud series temporales de " + var_periodo
    return descripcion


# Esta función construye el rango de fechas a partir de los inputs del usuario
def desc_datoscrud_cassandra(inicio, fin, sensor, codestacion):
    cur.execute(f''' select lm.station, s.name, lm.sensor, lm.event_time + interval '5' hour event_time, lm.event_value 
    from cassandra.raw.weather_events AS lm INNER JOIN cassandra.raw.stations AS s
        ON lm.station = s.stationid
    where lm.station in ('{codestacion}')
    AND lm.event_time BETWEEN timestamp '{inicio}' AND timestamp '{fin}' 
    AND lm.sensor in ('{sensor}') ''')  # estación de la forma '0021206600'; sensor de la forma '0240'
    station_sensem = cur.fetchall()

    # Se genera un Data Frame a partir de los datos extaídos
    stationdf = pd.DataFrame(station_sensem, columns=[
                             'Station', 'Name', 'Sensor', 'Fecha', 'Valor'])

    return station_sensem, stationdf[['Fecha', 'Valor']]


def reemplazar_datos_en_runs(paragraph, search_text, replace_text):
    """Busca y reemplaza texto en los runs de un párrafo, intentando mantener el estilo."""
    text = paragraph.text  # Obtener todo el texto del párrafo
    if search_text in text:
        # Reemplazar el texto a nivel de párrafo
        new_text = text.replace(search_text, replace_text)
        # Borrar el texto actual del párrafo
        for run in paragraph.runs:
            run.text = ''
            run.font.name = 'Verdana'
            run.font.size = Pt(11)
        # Asignar el nuevo texto al primer run para mantener el estilo inicial del párrafo
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:  # Si no hay runs, simplemente añadir el texto al párrafo
            paragraph.add_run(new_text)

def reemplazar_datos_noEMC(doc, nombres, apellidos, correo, descrip_solicit, lat_pi, lon_pi):
    #doc = Document(nombre_plantilla)
    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidos.upper(),
        "{{CORREO}}": correo,
        "{{DESCRIP_SOLICIT}}": descrip_solicit,
        "{{LAT_PI}}": str(lat_pi),
        "{{LONG_PI}}": str(lon_pi)
    }
    
    # Reemplazar datos en los runs de cada párrafo para mantener el estilo
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)

    return doc
    
def reemplazar_datos_nan(doc, nombres, apellidos, correo, dias, meses, ano, selected_variable, estacion_nombre, descrip_solicit):
    #doc = Document(nombre_plantilla)
    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidos.upper(),
        "{{CORREO}}": correo,
        "{{DIAS}}": ", ".join(map(str, dias)),
        "{{MESES}}": ", ".join(meses),
        "{{AÑO}}": ", ".join(map(str, ano)),
        "{{VARIABLE}}": selected_variable,
        "{{ESTACION}}": estacion_nombre,
        "{{DESCRIP_SOLICIT}}": descrip_solicit
    }

    # Reemplazar datos en los runs de cada párrafo para mantener el estilo
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)

    return doc


def reemplazar_datos_precip(nombre_plantilla, nombres, apellidos, dias, meses, ano, selected_variable, estacion_seleccionada, dia, mes_nm, ano_p, primer_valor):
    doc = Document(nombre_plantilla)
    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidos.upper(),
        "{{DIAS}}": ", ".join(map(str, dias)),
        "{{MESES}}": ", ".join(meses),
        "{{AÑO}}": ", ".join(map(str, ano)),
        "{{VARIABLE}}": selected_variable,
        "{{ESTACION}}": estacion_seleccionada['nombre'],
        "{{LATITUD}}": str(estacion_seleccionada['latitud']),
        "{{LONGITUD}}": str(estacion_seleccionada['longitud']),
        "{{ALTITUD}}": str(estacion_seleccionada['altitudDEM']),
        "{{MUNICIPIO}}": str(estacion_seleccionada['MUNICIPIO']),
        "{{DEPARTAMENTO}}": str(estacion_seleccionada['DEPARTAMENTO']),
        "{{DIA}}": str(dia),
        "{{MES}}": str(mes_nm),
        "{{AÑO_P}}": str(ano_p),
        "{{PRIMER_DATO}}": str(primer_valor),
        "{{PRIMER_DATO_HA}}": str(primer_valor * 10),
    }

    # Reemplazar datos en los runs de cada párrafo para mantener el estilo
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)

    return doc


def reemplazar_datos(nombre_plantilla, nombres, apellidos, dias, meses, ano, selected_variable, estacion_seleccionada):
    doc = Document(nombre_plantilla)
    datos = {
        "{{NOMBRES}}": nombres.upper(),
        "{{APELLIDOS}}": apellidos.upper(),
        "{{DIAS}}": ", ".join(map(str, dias)),
        "{{MESES}}": ", ".join(meses),
        "{{AÑO}}": ", ".join(map(str, ano)),
        "{{VARIABLE}}": selected_variable,
        "{{ESTACION}}": estacion_seleccionada['nombre'],
        "{{LATITUD}}": str(estacion_seleccionada['latitud']),
        "{{LONGITUD}}": str(estacion_seleccionada['longitud']),
        "{{ALTITUD}}": str(estacion_seleccionada['altitud']),
        "{{MUNICIPIO}}": str(estacion_seleccionada['MUNICIPIO']),
        "{{DEPARTAMENTO}}": str(estacion_seleccionada['DEPARTAMENTO']),
    }

    # Reemplazar datos en los runs de cada párrafo para mantener el estilo
    for p in doc.paragraphs:
        for key, value in datos.items():
            reemplazar_datos_en_runs(p, key, value)

    return doc


@app.callback(
    Output('distance-info', 'children'),
    Input('calculate-distance-btn', 'n_clicks'),
    State('mapa-estaciones', 'clickData'),
    State('lat-input', 'value'),
    State('lon-input', 'value')
)
def calculate_distance(n_clicks, clickData, lat, lon):
    if n_clicks is None or clickData is None:
        return html.Div("Seleccione una estación, luego, cuando vea el marcador rojo, introduzca las coordenadas \
            de su punto de interés y oprima el botón 'Calcular distancia'.", style={'font-family': 'arial', 'font-size': 13})

    estacion_lat = clickData['points'][0]['lat']
    estacion_lon = clickData['points'][0]['lon']

    if lat is None or lon is None:
        return html.Div("Por favor, introduzca las coordenadas del punto de interés.", style={'font-family': 'arial', 'font-size': 13})

    distancia = geopy.distance.distance(
        (estacion_lat, estacion_lon), (lat, lon)).km
    return html.Div(f"La distancia es: {distancia:.2f} km.", style={'font-family': 'arial', 'font-size': 15})

# Callback de actualizar mapa
@app.callback(
    [Output('mapa-estaciones', 'figure'),
     Output('lat-input', 'value'),
     Output('lon-input', 'value')],
    [Input('reset-button', 'n_clicks'),
     Input('mapa-estaciones', 'clickData')],
    prevent_initial_call=True
)
def update_map(n_clicks, clickData):
    ctx = dash.callback_context
    triggered_id = ctx.triggered[0]['prop_id'].split('.')[0]

    if triggered_id == 'reset-button':
        fig = px.scatter_mapbox(EstCON_Act,
                                lat="latitud",
                                lon="longitud",
                                hover_name="nombre",
                                zoom=10,
                                mapbox_style="open-street-map")
        return fig, None, None  # Reset everything

    if clickData:
        lat = clickData['points'][0]['lat']
        lon = clickData['points'][0]['lon']
        fig = px.scatter_mapbox(EstCON_Act,
                                lat="latitud",
                                lon="longitud",
                                hover_name="nombre",
                                zoom=10,
                                mapbox_style="open-street-map")
        fig.add_scattermapbox(lat=[lat], lon=[lon], marker={
                              'size': 15, 'color': 'red'})
        return fig, None, None  # Update map with the station marker

    return dash.no_update

# Callback de cuando no hay estaciones
@app.callback(
    Output('hidden-div', 'children'),  # Este es un componente oculto para almacenar los datos
    [Input('sin-estaciones', 'value'),
     State('lat-input', 'value'),
     State('lon-input', 'value')]
)
def handle_no_station(selection, lat, lon):
    if selection == 'no_station':
        return json.dumps({'lat': lat, 'lon': lon})
    return dash.no_update

# Callback para actualizar opciones segundo dropdown de variables
@app.callback(
    Output("tiposerie-dropdown", "options"),
    [Input("variable-dropdown", "value")]
)
def set_options(selected_variable):
    if selected_variable == "Precipitación":
        return [{"label": "Precipitación total diaria",
                 "value": "Precipitación total diaria"},
                {"label": "Precipitación total mensual",
                 "value": "Precipitación total mensual"},
                {"label": "Precipitación total anual",
                 "value": "Precipitación total anual"}]
    elif selected_variable == "Temperatura máxima":
        return [{"label": "Temperatura máxima diaria",
                 "value": "Temperatura máxima diaria"},
                {"label": "Temperatura máxima media mensual",
                 "value": "Temperatura máxima media mensual"},
                {"label": "Temperatura máxima media anual",
                 "value": "Temperatura máxima media anual"}]
    elif selected_variable == "Temperatura mínima":
        return [{"label": "Temperatura mínima diaria",
                 "value": "Temperatura mínima diaria"},
                {"label": "Temperatura mínima media mensual",
                 "value": "Temperatura mínima media mensual"},
                {"label": "Temperatura mínima media anual",
                 "value": "Temperatura mínima media anual"}]
    elif selected_variable == "Temperatura del aire":
        return [{"label": "Temperatura del aire media diaria",
                 "value": "Temperatura del aire media diaria"},
                {"label": "Temperatura del aire media mensual",
                 "value": "Temperatura del aire media mensual"},
                {"label": "Temperatura del aire media anual",
                 "value": "Temperatura del aire media anual"}]
    elif selected_variable == "Velocidad del viento":
        return [{"label": "Velocidad del viento horaria",
                 "value": "Velocidad del viento horaria"},
                {"label": "Velocidad del viento media diaria",
                 "value": "Velocidad del viento media diaria"},
                {"label": "Velocidad del viento media mensual",
                 "value": "Velocidad del viento media mensual"},
                {"label": "Velocidad del viento media anual",
                 "value": "Velocidad del viento media anual"}]
    return []


plantillas_por_variable = {
    "Precipitación total diaria": "PlantillaPrecipDiaria.docx",
    "Precipitación total mensual": "PlantillaPrecipMensual.docx",
    "Precipitación total anual": "PlantillaPrecipAnual.docx",
    "Temperatura máxima diaria": "PlantillaTemperatDiaria.docx",
    "Temperatura máxima media mensual": "PlantillaTemperatMensual.docx",
    "Temperatura máxima media anual": "PlantillaTemperatAnual.docx",
    "Temperatura mínima diaria": "PlantillaTemperatDiaria.docx",
    "Temperatura mínima media mensual": "PlantillaTemperatMensual.docx",
    "Temperatura mínima media anual": "PlantillaTemperatAnual.docx",
    "Temperatura del aire media diaria": "PlantillaTemperatDiaria.docx",
    "Temperatura del aire media mensual": "PlantillaTemperatMensual.docx",
    "Temperatura del aire media anual": "PlantillaTemperatAnual.docx",
    "Velocidad del viento horaria": "PlantillaVelVientoHoraria.docx",
    "Velocidad del viento media diaria": "PlantillaVelVientoDiaria.docx",
    "Velocidad del viento media mensual": "PlantillaVelVientoMensual.docx",
    "Velocidad del viento media anual": "PlantillaVelVientoAnual.docx",
    "Sin Datos": "PlantillaOficioLamentoSinDatos.docx",
    "Sin Estación": "PlantillaOficioLamentoSinEstaciones.docx"
}

# Callback para generar la certificación al hacer clic en el botón "generar certificado"
@app.callback(
    # Se genera el output completo
    # Asegurarse de que este Output coincida con un componente en tu layout
    Output("output-state", "children"),
    Input("generar-button", "n_clicks"),
    [State("nombres-input", "value"),
     State("apellidos-input", "value"),
     State("correo-input", "value"),
     State("dias-dropdown", "value"),
     State("meses-dropdown", "value"),
     State("ano-dropdown", "value"),
     State("tiposerie-dropdown", "value"),
     State("estacion-dropdown", "value"),
     State("sin-estaciones", "value"),
     State("lat-input", "value"),
     State("lon-input", "value")]
)
def generar_certificado(n_clicks, nombres, apellidos, correo, dias, meses, ano, selected_variable, estacion_nombre, sin_estacion, lat, lon):
    if n_clicks is not None:
        try:
            # Asegúrate de que los valores no sean None o vacíos.
            if not nombres or not apellidos or not dias or not meses or not ano or not selected_variable or not estacion_nombre:
                return html.Div("Por favor, diligencie completamente el formulario para obtener su certificación.",
                                style={'font-family': 'Arial', 'font-style': 'italic', 'color': 'darkred', 'font-size': 13})

            descrip_solicit = construir_descripsolicit(selected_variable)
            # Primero que todo, se hace el resultado de punto de interés sin estaciones
            if sin_estacion == 'no_station':
                # Se selecciona la plantilla para "Sin Estación"
                nombre_plantilla = plantillas_por_variable["Sin Estación"]
                doc = Document(nombre_plantilla)
                # Llama a tu función de reemplazo de datos especial para el caso sin estación
                doc = reemplazar_datos_noEMC(doc, nombres, apellidos, correo, descrip_solicit, lat, lon)
                # Guarda o procesa el documento como necesario
                nombre_archivo_final = f"Modif_{nombre_plantilla}"
                # Guarda el documento Word modificado
                doc.save(nombre_archivo_final)
                # Generar el output HTML o lo que sea necesario para mostrar al usuario
                return html.Div("Respuesta generada para punto de interés sin estaciones cercanas representativas.",
                                style={'font-family': 'Arial', 'font-style': 'italic', 'color': 'darkorange', 'font-size': 13})
            
            # Suponiendo que tienes un DataFrame que corresponde a la estación seleccionada
            estacion_seleccionada = EstCON_Act[EstCON_Act['nombre']
                                               == estacion_nombre].iloc[0]

            # Construir rango de fechas, obtener sensor, y construir código de estación
            inicio, fin = construir_rango_fechas(dias, meses, ano)
            sensor = obtener_sensor(selected_variable)
            codestacion = construir_codestacion(estacion_seleccionada)

            # Realizar consulta a la base de datos
            _, stationdf = desc_datoscrud_cassandra(
                inicio, fin, sensor, codestacion)
            
            # Si escoge viento, se hace la transformación a km/h de una vez
            if 'viento' in selected_variable:
                stationdf['Valor'] = (stationdf['Valor'] * 3.6).round(1)

            def aplicar_transformacion(df, sel_var):
                # Diccionario inicial para casos especiales
                rawdata_to_agg = {
                    "Precipitación total diaria": lambda df: df,
                    "Precipitación total mensual": lambda df: dvd.PT_10_TT_M(df),
                    "Precipitación total anual": lambda df: dvd.PT_10_TT_A(df),
                    "Temperatura del aire media diaria": lambda df: dvd.TSSM_MEDIA_D(df),
                    "Velocidad del viento media diaria": lambda df: dvd.VVAG_D(df)
                }

                # Intenta primero encontrar la función en el diccionario especial
                funcion = rawdata_to_agg.get(sel_var)

                # Si no se encontró en el diccionario especial, aplicar lógica basada en sufijos
                if funcion is None:
                    if "mensual" in sel_var:
                        def funcion(df): return dvd.T2M_VVDAG_MEDIA_M(df)
                    elif "anual" in sel_var:
                        def funcion(df): return dvd.T2M_VVDAG_MEDIA_A(df)
                    else:
                        # Por defecto, si no es mensual o anual, asumir que no necesita transformación
                        def funcion(df): return df

                # Aplica la función seleccionada
                stationdf_fnl = funcion(df)

                return stationdf_fnl

            stationdf_fnl = aplicar_transformacion(stationdf, selected_variable)
            modifdfprecip_ClasifLimSup(stationdf_fnl, EstCON_Act, selected_variable, codestacion)
            normales = pd.read_excel('NormClimatolEstándar_PRECIPITACION_9120.xlsx')

            # Crear una función para extraer el valor normal del mes correspondiente
            def get_normal_value(row, normales_df):
                # Usar el mes para seleccionar la columna correcta
                month_column = row['Mes']
                # Encontrar el valor normal para la estación y el mes específicos
                codestacion_n = re.sub("[^0-9]", "", codestacion)
                try:
                    codestacion_n = int(codestacion_n)
                except ValueError:
                    print("codestacion contiene caracteres no convertibles a int.")
                normal_value = normales_df.loc[normales_df['Station']
                                               == codestacion_n, month_column]
                return normal_value.values[0] if not normal_value.empty else "Sin dato de ref."

            if selected_variable == "Precipitación total mensual":
                stationdf_fnl['Mes'] = stationdf_fnl['Fecha'].dt.month
                # Aplicar la función para crear una nueva columna con el valor normal
                stationdf_fnl['Valor_Normal'] = stationdf_fnl.apply(get_normal_value, normales_df=normales, axis=1)
                # Calcular la nueva columna como solicitaste
                stationdf_fnl['Índice (%)'] = ((stationdf_fnl['Valor'] * 100) / stationdf_fnl['Valor_Normal']).round(1)
                stationdf_fnl.drop(columns=['Mes', 'Valor_Normal'], inplace=True)

            # Si la variable escogida fue la velocidad horaria, se descarga también dirección del viento
            if selected_variable == "Velocidad del viento horaria":
                _, stationdf_dv = desc_datoscrud_cassandra(
                    inicio, fin, 'DVAG_CON', codestacion)
                # Cambiar nombre de columna de datos
                stationdf_dv.rename(
                    columns={'Valor': 'Dirección del viento (°)'}, inplace=True)
                # Se unen los dataframes
                stationdf_fnl = pd.merge(
                    stationdf_fnl, stationdf_dv, on='Fecha')

            # Determina si el DataFrame está vacío para seleccionar la plantilla
            clave_plantilla = "Sin Datos" if stationdf_fnl.empty else selected_variable

            # Se selecciona la plantilla basada en la clave
            nombre_plantilla = plantillas_por_variable.get(
                clave_plantilla, "PlantillaPorDefecto.docx")

            # Ahora inicializa el documento una sola vez aquí
            doc = Document(nombre_plantilla)

            # Realiza el reemplazo de datos. Ajusta esta llamada a tus necesidades específicas.
            if clave_plantilla == "Sin Datos":
                doc = reemplazar_datos_nan(doc, nombres, apellidos, correo, dias,
                                           meses, ano, selected_variable, estacion_nombre, descrip_solicit)
            else:
                # Esta variable contiene la selección correcta que mapea al diccionario
                clave_plantilla = selected_variable

                # Para que el nombre del mes quede en letras y español, se ajusta el código
                # Para Windows, prueba 'spanish_spain' o 'es_ES'
                locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
                # Se convierte la columna 'Fecha' a tipo datetime
                stationdf_fnl['Fecha'] = pd.to_datetime(stationdf_fnl['Fecha'], format='%Y-%m-%d %H:%M:%S.%f')

                # Se cambia según la temporalidad escogida por el usuario
                sufij_periodo = {"horaria": '%Y-%m-%d %H:%M',
                                 "diaria": '%Y-%m-%d', "mensual": '%Y-%m', "anual": '%Y'}
                for sufij, period in sufij_periodo.items():
                    if selected_variable.endswith(sufij):
                        stationdf_fnl['Fecha'] = stationdf_fnl['Fecha'].dt.strftime(
                            period)

                if clave_plantilla == "Precipitación total diaria":
                    primer_fecha = pd.to_datetime(stationdf_fnl['Fecha'].iloc[0])
                    dia = primer_fecha.day
                    # Para que el nombre del mes quede en letras y español, se ajusta según Windows
                    locale.setlocale(locale.LC_TIME, 'es_ES.UTF-8')
                    # Se tiene el nombre del mes en español
                    mes_nm = primer_fecha.strftime('%B')
                    ano_p = primer_fecha.year
                    # Se extrae el primer valor
                    primer_valor = stationdf_fnl['Valor'].iloc[0]

                    doc = reemplazar_datos_precip(nombre_plantilla, nombres, apellidos, dias, meses,
                                                  ano, selected_variable, estacion_seleccionada, dia, mes_nm, ano_p, primer_valor)
                else:
                    # Luego de tener todos los datos necesarios, llamar a reemplazar_datos
                    doc = reemplazar_datos(nombre_plantilla, nombres, apellidos,
                                           dias, meses, ano, selected_variable, estacion_seleccionada)

                und = set_und(selected_variable)

                # Se llenan valores vacíos con 'ND' según se tiene estipulado
                stationdf_fnl['Valor'].fillna('ND', inplace=True)
                # Cambiar nombre de columna de datos
                stationdf_fnl.rename(columns={'Valor': selected_variable + f" ({und})"}, inplace=True)

                table = doc.add_table(
                    rows=stationdf_fnl.shape[0] + 1, cols=stationdf_fnl.shape[1])
                table.style = 'Table Grid'
                for i, column in enumerate(stationdf_fnl.columns):
                    table.cell(0, i).text = column
                    for j, value in enumerate(stationdf_fnl[column]):
                        table.cell(j + 1, i).text = str(value)

                # Encuentra el párrafo específico en el documento donde deseas insertar la tabla
                def move_table_after(table, paragraph):
                    tbl, p = table._tbl, paragraph._p
                    p.addnext(tbl)

                # Reemplaza con el texto de tu marcador
                marcador_texto = "Datos disponibles"
                for paragraph in doc.paragraphs:
                    if marcador_texto in paragraph.text:
                        # Inserta la tabla desde el archivo Excel después del párrafo encontrado
                        # Agrega un salto de línea antes de la tabla
                        doc.add_paragraph('\n')
                        move_table_after(table, paragraph)
                        break  # Sal del bucle después de insertar la tabla  # Sal del bucle después de insertar la tabla

                # Configurar la primera fila para repetirse en cada nueva página
                row = table.rows[0]
                tr = row._tr
                tblPr = tr.get_or_add_trPr()
                tblHeader = OxmlElement('w:tblHeader')
                tblHeader.set(qn('w:val'), "true")
                tblPr.append(tblHeader)

            nombre_archivo_final = f"Modif_{nombre_plantilla}"
            # Guarda el documento Word modificado
            doc.save(nombre_archivo_final)

            return html.Div("Se generó la certificación.", style={'font-family': 'Arial', 'font-style': 'italic', 'color': 'darkgreen', 'font-size': 13})
        except Exception as e:
            # return html.Div(f"Se produjo un error al generar la certificación: {e}", style={'font-family': 'Arial', 'font-style': 'italic', 'color':'red', 'font-size':13})
            # Captura el traceback completo
            error_traceback = traceback.format_exc()
            # Retorna un mensaje de error más detallado incluyendo el traceback
            return html.Div([html.Div(f"Intente más tarde, se produjo un error al generar la certificación: {e}",
                                      style={'font-family': 'Arial', 'font-style': 'italic', 'color': 'red', 'font-size': 13}),
                             html.Pre(error_traceback,
                                      style={'font-family': 'Consolas', 'font-style': 'italic', 'color': 'grey', 'font-size': 10})])
    return html.Div("Haga clic en este botón para generar la certificación:", style={'font-family': 'Arial', 'font-style': 'italic', 'font-weight': 'bold', 'font-size': 13})


if __name__ == "__main__":
    app.run_server(debug=True)
